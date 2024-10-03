from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph


# Create a new docx document
doc = Document()


def parse_style_attribute(style_attr):
    styles = {}
    for item in style_attr.split(';'):
        if ':' in item:
            key, value = item.split(':', 1)
            key = key.strip()
            value = value.strip()
            if key == 'font-size':
                if value.endswith('pt'):
                    size = float(value[:-2])
                else:
                    size = float(value)
                styles['font_size'] = size
            elif key == 'font-family':
                styles['font_name'] = value
            elif key == 'color':
                styles['font_color'] = value.lstrip('#')
            elif key == 'text-align':
                styles['text-align'] = value
            elif key == 'font-weight':
                if value == 'bold':
                    styles['bold'] = True
            elif key == 'font-style':
                if value == 'italic':
                    styles['italic'] = True
            elif key == 'padding-left' or key == 'margin-left':
                if value.endswith('px'):
                    indent = float(value[:-2])
                    styles['left_indent'] = Pt(indent * 0.75)
                elif value.endswith('pt'):
                    indent = float(value[:-2])
                    styles['left_indent'] = Pt(indent)
    return styles

def process_element(element, parent=None, style_context=None, list_style=None):
    if style_context is None:
        style_context = {}

    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            if isinstance(parent, _Cell):
                p = parent.add_paragraph()
                run = p.add_run(text)
            elif isinstance(parent, Paragraph):
                run = parent.add_run(text)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
            run.bold = style_context.get('bold', False)
            run.italic = style_context.get('italic', False)
            font_name = style_context.get('font_name', None)
            font_size = style_context.get('font_size', None)
            font_color = style_context.get('font_color', None)
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if font_size:
                run.font.size = Pt(font_size)
            if font_color:
                run.font.color.rgb = RGBColor.from_string(font_color)
    elif isinstance(element, Tag):
        style_attr = element.get('style', '')
        new_style_context = style_context.copy()
        styles = parse_style_attribute(style_attr)
        new_style_context.update(styles)

        if element.name in ['p', 'div']:
            if isinstance(parent, _Cell):
                p = parent.add_paragraph()
            else:
                p = doc.add_paragraph()
            if 'text-align' in styles:
                alignment = styles['text-align']
                if alignment == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if 'left_indent' in styles:
                p.paragraph_format.left_indent = styles['left_indent']
            for child in element.contents:
                process_element(child, parent=p, style_context=new_style_context)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                process_element(li, parent=parent, style_context=style_context, list_style=element.name)
        elif element.name == 'li':
            if isinstance(parent, _Cell):
                p = parent.add_paragraph()
            else:
                p = doc.add_paragraph()
            if list_style == 'ul':
                p.style = 'List Bullet'
            elif list_style == 'ol':
                p.style = 'List Number'
            for child in element.contents:
                process_element(child, parent=p, style_context=style_context)
        elif element.name == 'span':
            for child in element.contents:
                process_element(child, parent=parent, style_context=new_style_context)
        elif element.name == 'strong':
            new_style_context['bold'] = True
            for child in element.contents:
                process_element(child, parent=parent, style_context=new_style_context)
        elif element.name == 'em':
            new_style_context['italic'] = True
            for child in element.contents:
                process_element(child, parent=parent, style_context=new_style_context)
        elif element.name == 'br':
            if parent is not None and isinstance(parent, Paragraph):
                parent.add_run().add_break()
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'

                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    row_cells = table.add_row().cells
                    for j, cell_element in enumerate(cells):
                        if j >= len(row_cells):
                            table.add_column()
                            row_cells = table.rows[-1].cells
                        cell = row_cells[j]
                        process_element(cell_element, parent=cell, style_context=style_context)
                doc.add_paragraph('')
        else:
            for child in element.contents:
                process_element(child, parent=parent, style_context=new_style_context)


def docxifier(input_html, output_docx):



    # Load the HTML content
    with open(input_html, 'r', encoding='utf-8') as file:
        html_content = file.read()

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    # Set default font style for the whole document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Ensure the font is applied to all paragraph styles
    for s in doc.styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            s.font.name = 'Calibri'
            s.font.size = Pt(11)

    # Process the body of the HTML document
    body = soup.body
    for element in body.contents:
        process_element(element)

    # Save the document
    
    doc.save(output_docx)

    print("Document created successfully.")
