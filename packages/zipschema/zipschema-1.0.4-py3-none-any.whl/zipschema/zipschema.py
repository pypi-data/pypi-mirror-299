import os
import yaml
import zipfile
import json
import jsonschema
import jinja2
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from argparse import ArgumentParser


# Helper function to create a file tree
def generate_file_tree(schema_data):
    file_paths = []

    # Collect all file paths
    for element in schema_data.get('elements', []):
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed']:
            if key in element:
                for item in element[key]:
                    file_paths.append(item['path'])

    # Create file tree structure (basic indentation)
    file_tree = {}
    for path in file_paths:
        parts = path.split('/')
        current = file_tree
        total = len(parts)
        for i, part in enumerate(parts):
            if total > 1:
                if i+1 < total:
                    current = current.setdefault(f"{part}/", {})
                else:
                    current = current.setdefault(f"{part}", {})

    # Convert file tree to string representation
    def tree_to_string(tree, indent=0):
        result = ''
        for key, subtree in tree.items():
            result += '  ' * indent + f"- {key}\n"
            result += tree_to_string(subtree, indent + 1)
        return result

    return tree_to_string(file_tree)


# Validate the schema itself
def validate_schema(schema_data):
    required_keys = ['name', 'description', 'version', 'elements']
    for key in required_keys:
        if key not in schema_data:
            return False, f"Missing required field: {key}"

    return True, "Schema is valid."


# Helper function to validate a file against a JSON schema
def validate_jsonschema_file(zip_file, json_schema_path, file_path):
    with zip_file.open(file_path) as json_file:
        json_data = json.load(json_file)
        with open(json_schema_path, 'r') as schema_file:
            json_schema = json.load(schema_file)
            try:
                jsonschema.validate(instance=json_data, schema=json_schema)
                return True, f"{file_path} is valid according to the JSON schema."
            except jsonschema.exceptions.ValidationError as e:
                return False, f"Validation failed for {file_path}: {str(e)}"


# Helper function to recursively validate using zipschema
def validate_zipschema(zip_file, zipschema_path, file_path):
    with open(zipschema_path, 'r') as schema_file:
        zipschema = yaml.safe_load(schema_file)
    # Recursively validate the zip file using the zipschema
    result, message = validate_zip_against_schema(zip_file.filename, zipschema)
    if result:
        return True, f"{file_path} validated successfully using zipschema."
    else:
        return False, f"Validation failed for {file_path} using zipschema: {message}"


# Validate a zip file against the schema
def validate_zip_against_schema(zip_path, schema_data):
    if not zipfile.is_zipfile(zip_path):
        return False, "Provided file is not a valid zip file."

    with zipfile.ZipFile(zip_path, 'r') as zfile:
        zip_contents = zfile.namelist()

        for element in schema_data.get('elements', []):
            for key in ['allOf', 'anyOf', 'oneOf']:
                if key in element:
                    one_of_found = 0
                    for item in element[key]:
                        if item['path'] in zip_contents:
                            # If the item has a schema field
                            if 'schema' in item:
                                schema_type = item['schema'].split('.')[-1].lower()

                                # JSON schema validation
                                if schema_type == 'jsonschema':
                                    result, message = validate_jsonschema_file(zfile, item['schema'], item['path'])
                                    if not result:
                                        return False, message

                                # zipschema recursive validation
                                elif schema_type == 'zipschema' or item['schema'] == 'self':
                                    if item['schema'] == 'self':
                                        result, message = validate_zip_against_schema(zip_path, schema_data)
                                    else:
                                        result, message = validate_zipschema(zfile, item['schema'], item['path'])
                                    if not result:
                                        return False, message

                            # Validation for oneOf
                            if key == 'oneOf':
                                one_of_found += 1

                    if key == 'oneOf' and one_of_found != 1:
                        return False, f"Exactly one file from 'oneOf' should be present in section: {element['section_title']}"

            if 'noneOf' in element:
                for prohibited_file in element['noneOf']:
                    for zip_item in zip_contents:
                        if zip_item.startswith(prohibited_file['path'].split('*')[0]):
                            return False, f"Prohibited file found: {zip_item}"

    return True, "Zip file contents are valid."


# Function to add borders to DOCX table
def add_borders_to_cells(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'start', 'bottom', 'end']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Border size
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)


# Helper function to handle description output for Markdown
def format_description_markdown(description):
    if isinstance(description, list):
        return "\n\n".join(description)  # Each string in the list on a new line
    return description


# Helper function to handle description output for DOCX
def format_description_docx(doc, description):
    if isinstance(description, list):
        for line in description:
            doc.add_paragraph(line)  # Each string in the list as a separate paragraph
    else:
        doc.add_paragraph(description)


# Generate Markdown documentation with file table, section/conditional headings, and file items
def generate_markdown_with_tree(schema_data, output_file):
    file_tree = generate_file_tree(schema_data)

    template = """
# {{ schema.name }}

**Version:** {{ schema.version }}

{{ schema.description }}

## File Tree
{{ file_tree }}

## File List
| Filename | Summary | Section |
|----------|---------|---------|
{% for element in schema.elements %}
{% for key, items in element.items() if key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed'] %}
{% for item in items %}
| `{{ item.path }}` | {{ item.summary or '' }} | {{ element.section_title }} |
{% endfor %}
{% endfor %}
{% endfor %}

## Section List
{% for element in schema.elements %}
### {{ element.section_title }}
{{ format_description_markdown(element.section_description) }}

{% for key, items in element.items() if key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed'] %}
#### {{ key }}
{% for item in items %}
- **`{{ item.path }}`**: {{ item.description }} 
{% endfor %}
{% endfor %}
{% endfor %}
"""

    env = jinja2.Environment(loader=jinja2.BaseLoader, trim_blocks=True, lstrip_blocks=True)
    template = env.from_string(template)
    if "description" in schema_data:
        schema_data["description"] = format_description_markdown(schema_data["description"])
    markdown_output = template.render(schema=schema_data, file_tree=file_tree,
                                      format_description_markdown=format_description_markdown)

    with open(output_file, 'w') as md_file:
        md_file.write(markdown_output)

    return "Markdown documentation with file table, sections, conditionals, and file items generated."


# Generate DOCX documentation with file table, section/conditional headings, and file items
def generate_docx_with_tree(schema_data, output_file):
    doc = docx.Document()

    # Title and version
    doc.add_heading(schema_data['name'], 0)
    doc.add_paragraph(f"Version: {schema_data['version']}")

    # Description
    format_description_docx(doc, schema_data['description'])

    # File Tree Section
    doc.add_heading('File Tree', level=1)
    file_tree = generate_file_tree(schema_data)
    doc.add_paragraph(file_tree)

    # File List Section (Table)
    doc.add_heading('File List', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filename'
    hdr_cells[1].text = 'Summary'
    hdr_cells[2].text = 'Section'

    for element in schema_data.get('elements', []):
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed']:
            if key in element:
                for item in element[key]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item['path']
                    row_cells[1].text = item.get('summary', '')
                    row_cells[2].text = element['section_title']

    # Add borders to table
    add_borders_to_cells(table)

    # Section List
    for element in schema_data.get('elements', []):
        # Section Title
        doc.add_heading(element['section_title'], level=2)
        format_description_docx(doc, element['section_description'])

        # Conditional sub-sections (allOf, anyOf, oneOf, noneOf, allowed)
        for key in ['allOf', 'anyOf', 'oneOf', 'noneOf', 'allowed']:
            if key in element:
                doc.add_heading(f"{key}:", level=3)
                for item in element[key]:
                    doc.add_paragraph(f"File: {item['path']}\nDescription: {item['description']}")

    doc.save(output_file)

    return "DOCX documentation with file table, sections, conditionals, and file items generated."


# Command-line interface logic
def main():
    parser = ArgumentParser(description="Schema validator and zip file validator with documentation generator.")
    parser.add_argument('mode', choices=['validate-schema', 'validate-zip', 'generate-doc'],
                        help="The mode to run the tool in.")
    parser.add_argument('schema_file', help="Path to the zipschema YAML file.")
    parser.add_argument('--zipfile', help="Path to the zip file (required for zip validation).")
    parser.add_argument('--output', help="Path for the documentation output (optional for doc generation).")
    parser.add_argument('--format', choices=['markdown', 'docx'], default="markdown", help="Documentation format (markdown or docx).")

    args = parser.parse_args()

    # Load the schema file
    with open(args.schema_file, 'r') as schema_file:
        schema_data = yaml.safe_load(schema_file)

    # Handle the different modes
    if args.mode == 'validate-schema':
        valid, message = validate_schema(schema_data)
        print(message)

    elif args.mode == 'validate-zip':
        if not args.zipfile:
            print("Zip file path is required for zip validation.")
            return
        valid, message = validate_zip_against_schema(args.zipfile, schema_data)
        print(message)

    elif args.mode == 'generate-doc':
        if not args.format:
            print("Format is required for documentation generation.")
            return

        # Determine the output file location if not provided
        if not args.output:
            schema_dir = os.path.dirname(args.schema_file)
            schema_base_name = os.path.splitext(os.path.basename(args.schema_file))[0]
            if args.format == 'markdown':
                args.output = os.path.join(schema_dir, f"{schema_base_name}.md")
            elif args.format == 'docx':
                args.output = os.path.join(schema_dir, f"{schema_base_name}.docx")

        # Generate the appropriate documentation format
        if args.format == 'markdown':
            message = generate_markdown_with_tree(schema_data, args.output)
        elif args.format == 'docx':
            message = generate_docx_with_tree(schema_data, args.output)
        print(message)


if __name__ == '__main__':
    main()
